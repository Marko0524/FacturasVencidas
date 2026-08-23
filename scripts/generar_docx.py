"""Genera docs/preguntas-tecnicas.docx a partir del markdown.

El markdown es la fuente de verdad: el .docx se regenera, no se edita a mano.
Así el documento de Word nunca se desincroniza del repositorio.

    pip install python-docx
    python scripts/generar_docx.py
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
SOURCE = PROJECT_ROOT / "docs" / "preguntas-tecnicas.md"
TARGET = PROJECT_ROOT / "docs" / "preguntas-tecnicas.docx"

# Misma paleta que los diagramas del asistente, para que el paquete se vea de una pieza.
TEAL = RGBColor(0x0E, 0x6E, 0x7A)
INK = RGBColor(0x14, 0x20, 0x2B)
SLATE = RGBColor(0x4A, 0x5A, 0x68)
CALLOUT_BG = "EDF3F4"
HEADER_BG = "E7EDF1"

BODY_FONT = "Georgia"
HEAD_FONT = "Segoe UI Semibold"
MONO_FONT = "Consolas"

# Dos pasadas, porque el markdown anida: **negrita con *cursiva* dentro**.
# Una sola expresión plana dejaría los asteriscos literales en el documento.
BOLD = re.compile(r"\*\*(.+?)\*\*", re.S)
INLINE = re.compile(
    r"\*(?P<italic>[^*]+)\*"
    r"|`(?P<code>[^`]+)`"
    r"|\[(?P<text>[^\]]+)\]\((?P<href>[^)]+)\)"
)


def git_user_name() -> str:
    """Autor del documento: el mismo que firma los commits."""
    try:
        out = subprocess.run(
            ["git", "config", "user.name"],
            cwd=PROJECT_ROOT, capture_output=True, text=True, timeout=10, check=False,
        )
        return out.stdout.strip() or "Autor"
    except (OSError, subprocess.SubprocessError):
        return "Autor"


def _shd(fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    return shd


def shade_paragraph(paragraph, fill: str) -> None:
    paragraph._p.get_or_add_pPr().append(_shd(fill))


def shade_cell(cell, fill: str) -> None:
    cell._tc.get_or_add_tcPr().append(_shd(fill))


def add_runs(paragraph, text: str) -> None:
    """Escribe el texto resolviendo negritas, cursivas, código y enlaces."""
    pos = 0
    for match in BOLD.finditer(text):
        if match.start() > pos:
            _add_inline(paragraph, text[pos:match.start()], bold=False)
        _add_inline(paragraph, match.group(1), bold=True)
        pos = match.end()
    if pos < len(text):
        _add_inline(paragraph, text[pos:], bold=False)


def _add_inline(paragraph, text: str, bold: bool) -> None:
    """Segunda pasada: cursiva, código y enlaces, heredando la negrita."""
    def emit(content: str, *, italic: bool = False, mono: bool = False):
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = MONO_FONT
            run.font.size = Pt(9.5)
        return run

    pos = 0
    for match in INLINE.finditer(text):
        if match.start() > pos:
            emit(text[pos:match.start()])
        if match.group("italic"):
            emit(match.group("italic"), italic=True)
        elif match.group("code"):
            emit(match.group("code"), mono=True)
        else:
            # Los enlaces relativos del repo no sirven como hipervínculo en Word:
            # se conserva el texto, que es la ruta del archivo.
            emit(match.group("text"), mono=True)
        pos = match.end()
    if pos < len(text):
        emit(text[pos:])


def style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    # Word necesita el nombre también en rFonts para que no caiga a la fuente por defecto.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)

    for level, size in ((1, 15), (2, 12)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = TEAL
        st.font.bold = True
        st.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        st.paragraph_format.space_after = Pt(6)
        rf = st.element.get_or_add_rPr().get_or_add_rFonts()
        rf.set(qn("w:ascii"), HEAD_FONT)
        rf.set(qn("w:hAnsi"), HEAD_FONT)


def add_page_numbers(doc: Document) -> None:
    """Pie de página con 'Página N' mediante un campo de Word."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Página ").font.size = Pt(8)

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    rpr.append(sz)
    run.append(rpr)
    fld.append(run)
    p._p.append(fld)

    for r in p.runs:
        r.font.color.rgb = SLATE
        r.font.name = BODY_FONT


def add_callout(doc: Document, lines: list[str]) -> None:
    """Bloque '> ...' del markdown, como recuadro sombreado."""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.right_indent = Inches(0.15)
        pf.space_before = Pt(6 if i == 0 else 0)
        pf.space_after = Pt(6 if i == len(lines) - 1 else 2)
        shade_paragraph(p, CALLOUT_BG)
        add_runs(p, line)
        for run in p.runs:
            run.font.size = Pt(9.5)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, *body = rows
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for col, text in enumerate(header):
        cell = table.cell(0, col)
        cell.text = ""
        shade_cell(cell, HEADER_BG)
        p = cell.paragraphs[0]
        add_runs(p, text)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = HEAD_FONT

    for r, row in enumerate(body, start=1):
        for col, text in enumerate(row):
            cell = table.cell(r, col)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if re.fullmatch(r"\s*\|[\s:|-]+\|\s*", line):
            continue  # separador de encabezado
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def build() -> None:
    doc = Document()
    style_base(doc)
    add_page_numbers(doc)

    author = git_user_name()
    doc.core_properties.title = "Preguntas técnicas conceptuales"
    doc.core_properties.author = author
    doc.core_properties.comments = "Generado desde docs/preguntas-tecnicas.md"

    lines = SOURCE.read_text(encoding="utf-8").split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            title = doc.add_paragraph()
            run = title.add_run(stripped[2:])
            run.font.name = HEAD_FONT
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = INK
            title.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                content = lines[i].strip().lstrip(">").strip()
                if content:
                    block.append(content)
                i += 1
            add_callout(doc, block)
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table(block))
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, lines[i].strip()[2:])
                i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                p = doc.add_paragraph(style="List Number")
                add_runs(p, re.sub(r"^\d+\.\s", "", lines[i].strip()))
                i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(TARGET)
    print(f"escrito {TARGET.relative_to(PROJECT_ROOT)} (autor: {author})")


if __name__ == "__main__":
    build()
